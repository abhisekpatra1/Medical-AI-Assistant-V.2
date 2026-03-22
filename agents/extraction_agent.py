"""
Extraction Agent
Extracts specific data, tables, and images from documents
"""

from typing import Dict, List
import os
from loguru import logger
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from pdf2image import convert_from_path
from PIL import Image


class ExtractionAgent:
    """
    Specialized agent for extracting exact content from documents
    Preserves tables, charts, and formatting
    """
    
    def __init__(self):
        logger.info("Extraction Agent initialized")
    
    async def extract_for_section(self, file_paths: List[str], section_name: str) -> Dict:
        """
        Extract content relevant to a specific report section
        """
        try:
            extracted = {
                "text": [],
                "tables": [],
                "images": []
            }
            
            for file_path in file_paths:
                ext = os.path.splitext(file_path)[1].lower()
                
                if ext == '.pdf':
                    content = self._extract_from_pdf(file_path, section_name)
                elif ext in ['.docx', '.doc']:
                    content = self._extract_from_docx(file_path, section_name)
                elif ext in ['.xlsx', '.xls']:
                    content = self._extract_from_excel(file_path, section_name)
                else:
                    continue
                
                # Merge extracted content
                extracted["text"].extend(content.get("text", []))
                extracted["tables"].extend(content.get("tables", []))
                extracted["images"].extend(content.get("images", []))
            
            logger.info(f"Extracted content for section: {section_name}")
            return extracted
        
        except Exception as e:
            logger.error(f"Error in extraction: {str(e)}")
            raise

    def _is_relevant_to_section(self, text: str, section_name: str) -> bool:
        """Check if text is relevant to the section"""
        text_lower = text.lower()
        section_lower = section_name.lower()
        
        section_keywords = {
            "introduction": ["background", "overview", "introduction", "purpose"],
            "clinical findings": ["findings", "results", "observations", "clinical", "diagnosis"],
            "patient tables": ["patient", "table", "data", "demographics"],
            "graphs": ["graph", "chart", "figure", "plot"],
            "summary": ["summary", "conclusion", "findings", "recommendations"]
        }
        
        keywords = section_keywords.get(section_lower, [])
        return any(keyword in text_lower for keyword in keywords) or section_lower in text_lower
    
    def _format_extracted_content(self, extracted: Dict) -> str:
        """Format extracted content for display"""
        formatted = []
        
        if extracted["text"]:
            formatted.append("Extracted Text:")
            formatted.extend(extracted["text"])
        
        if extracted["tables"]:
            formatted.append("\nExtracted Tables:")
            for i, table in enumerate(extracted["tables"], 1):
                formatted.append(f"\nTable {i}:")
                for row in table["data"]:
                    formatted.append(" | ".join(str(cell) for cell in row))
        
        if extracted["images"]:
            formatted.append(f"\nExtracted {len(extracted['images'])} image(s)")
        
        return "\n".join(formatted) if formatted else "No relevant content found."

    async def extract_from_query(self, file_paths: List[str], query: str) -> Dict:
        """
        Extract content based on specific query
        """
        extract_tables = any(word in query.lower() for word in ["table", "data", "rows", "columns"])
        extract_images = any(word in query.lower() for word in ["image", "chart", "graph", "figure"])
        
        extracted = {
            "text": [],
            "tables": [],
            "images": []
        }
        
        for file_path in file_paths:
            if extract_tables:
                tables = self._extract_all_tables(file_path)
                extracted["tables"].extend(tables)
            
            if extract_images:
                images = self._extract_all_images(file_path)
                extracted["images"].extend(images)
        
        content = self._format_extracted_content(extracted)
        sources = [os.path.basename(fp) for fp in file_paths]
        
        return {"content": content, "sources": sources}
    
    def _extract_from_pdf(self, file_path: str, section_name: str) -> Dict:
        """Extract content from PDF"""
        content = {"text": [], "tables": [], "images": []}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and self._is_relevant_to_section(text, section_name):
                        content["text"].append(text)
                    
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            content["tables"].append({"data": table, "source": file_path})
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
        
        return content
    
    def _extract_from_docx(self, file_path: str, section_name: str) -> Dict:
        """Extract content from Word document"""
        content = {"text": [], "tables": [], "images": []}
        
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text.strip() and self._is_relevant_to_section(para.text, section_name):
                    content["text"].append(para.text)
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append({"data": table_data, "source": file_path})
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
        
        return content
    
    def _extract_from_excel(self, file_path: str, section_name: str) -> Dict:
        """Extract content from Excel"""
        content = {"text": [], "tables": [], "images": []}
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else "" for cell in row])
                if data:
                    content["tables"].append({
                        "data": data,
                        "source": f"{file_path} - {sheet_name}"
                    })
        except Exception as e:
            logger.error(f"Error extracting from Excel: {str(e)}")
        
        return content
    
    def _extract_all_tables(self, file_path: str) -> List[Dict]:
        """Extract all tables from a document"""
        tables = []
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_tables = page.extract_tables()
                        for table in page_tables:
                            if table:
                                tables.append({"data": table, "source": file_path})
            elif ext in ['.docx', '.doc']:
                doc = DocxDocument(file_path)
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append({"data": table_data, "source": file_path})
            elif ext in ['.xlsx', '.xls']:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    data = []
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            data.append([str(cell) if cell is not None else "" for cell in row])
                    if data:
                        tables.append({"data": data, "source": f"{file_path} - {sheet_name}"})
        except Exception as e:
            logger.error(f"Error extracting tables: {str(e)}")
        
        return tables
    
    def _extract_all_images(self, file_path: str) -> List[Dict]:
        """Extract all images from a PDF document"""
        images = []
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                pdf_images = convert_from_path(file_path)
                for i, img in enumerate(pdf_images):
                    img_path = f"temp_img_{i}.png"
                    img.save(img_path)
                    images.append({"path": img_path, "source": file_path, "page": i + 1})
        except Exception as e:
            logger.error(f"Error extracting images: {str(e)}")
        
        return images
