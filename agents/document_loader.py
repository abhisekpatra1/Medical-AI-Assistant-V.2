"""
Document Loader Agent
Handles ingestion of multi-format medical documents
"""

from typing import List, Dict
import os
from loguru import logger
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract
import pdfplumber

from services.vector_store import VectorStoreService


class DocumentLoaderAgent:
    """
    Specialized agent for loading and processing various document formats
    Supports: PDF, DOCX, XLSX, Images (with OCR)
    """
    
    def __init__(self, vector_store: VectorStoreService):
        self.vector_store = vector_store
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        logger.info("Document Loader Agent initialized")
    
    async def load_and_process(self, session_id: str, file_paths: List[str]) -> str:
        """
        Load and process all documents
        
        Args:
            session_id: Session identifier
            file_paths: List of file paths to process
            
        Returns:
            Status message
        """
        try:
            all_documents = []
            
            for file_path in file_paths:
                logger.info(f"Processing file: {file_path}")
                
                # Determine file type and process accordingly
                ext = os.path.splitext(file_path)[1].lower()
                
                if ext == '.pdf':
                    docs = self._process_pdf(file_path)
                elif ext in ['.docx', '.doc']:
                    docs = self._process_docx(file_path)
                elif ext in ['.xlsx', '.xls']:
                    docs = self._process_excel(file_path)
                elif ext in ['.png', '.jpg', '.jpeg']:
                    docs = self._process_image(file_path)
                else:
                    logger.warning(f"Unsupported file type: {ext}")
                    continue
                
                all_documents.extend(docs)
            
            # Split documents into chunks
            chunks = self.text_splitter.split_documents(all_documents)
            
            # Store in vector database
            self.vector_store.add_documents(session_id, chunks)
            
            logger.info(f"Processed {len(all_documents)} documents into {len(chunks)} chunks")
            return f"Successfully processed {len(file_paths)} files"
            
        except Exception as e:
            logger.error(f"Error in document loading: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: str) -> List[Document]:
        """Extract text from PDF using pdfplumber for better table support"""
        documents = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    
                    if text:
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": i + 1,
                                "type": "pdf"
                            }
                        ))
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = self._format_table(table)
                            documents.append(Document(
                                page_content=table_text,
                                metadata={
                                    "source": file_path,
                                    "page": i + 1,
                                    "type": "pdf_table",
                                    "table_index": table_idx
                                }
                            ))
            
            logger.info(f"Extracted {len(documents)} sections from PDF")
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            # Fallback to PyPDF2
            documents = self._process_pdf_fallback(file_path)
        
        return documents
    
    def _process_pdf_fallback(self, file_path: str) -> List[Document]:
        """Fallback PDF processing using PyPDF2"""
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    if text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": i + 1,
                                "type": "pdf"
                            }
                        ))
        except Exception as e:
            logger.error(f"Fallback PDF processing failed: {str(e)}")
        
        return documents
    
    def _process_docx(self, file_path: str) -> List[Document]:
        """Extract text from Word documents"""
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            if full_text:
                documents.append(Document(
                    page_content="\n".join(full_text),
                    metadata={
                        "source": file_path,
                        "type": "docx"
                    }
                ))
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                table_text = self._format_table(table_data)
                documents.append(Document(
                    page_content=table_text,
                    metadata={
                        "source": file_path,
                        "type": "docx_table",
                        "table_index": table_idx
                    }
                ))
            
            logger.info(f"Extracted {len(documents)} sections from DOCX")
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
        
        return documents
    
    def _process_excel(self, file_path: str) -> List[Document]:
        """Extract data from Excel spreadsheets"""
        documents = []
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract all data
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else "" for cell in row])
                
                if data:
                    table_text = self._format_table(data)
                    documents.append(Document(
                        page_content=table_text,
                        metadata={
                            "source": file_path,
                            "sheet": sheet_name,
                            "type": "excel"
                        }
                    ))
            
            logger.info(f"Extracted {len(documents)} sheets from Excel")
            
        except Exception as e:
            logger.error(f"Error processing Excel: {str(e)}")
        
        return documents
    
    def _process_image(self, file_path: str) -> List[Document]:
        """Extract text from images using OCR"""
        documents = []
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "type": "image_ocr"
                    }
                ))
                logger.info(f"Extracted text from image via OCR")
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
        
        return documents
    
    def _format_table(self, table_data: List[List[str]]) -> str:
        """Format table data as text"""
        if not table_data:
            return ""
        
        # Create a formatted table string
        formatted = "TABLE:\n"
        for row in table_data:
            formatted += " | ".join(str(cell) for cell in row) + "\n"
        
        return formatted